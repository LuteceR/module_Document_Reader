import spacy
from spacy.matcher import Matcher
import docx
import pathlib
import mpire
import os
from mdr.config.students_keyWords import STUDENTS_IDENTIFICATIONS_WORDS
from mdr.config.teachers_keyWords import TEACHERS_IDENTIFICATIONS_WORDS
# import typing

nlp = spacy.load("ru_core_news_lg")

# def extract_names(text):
#     doc = nlp(text)
#     return [ent.text for ent in doc.ents if ent.label_ == "PER"]

# def lookingup(parallel: bool, array: list):
#     """
#     поиск статуса ФИО\n
#     parallel[bool] - распараллеливание поиска\n
#     array[list] - лист с нормализованным текстом\n
#     id_fio[int] - id расположения фио 
#     """
#     if not parallel:

class mdr:
    def __init__(self):
        
        """
        Создание объекта Module Document Reader
        """

        self.path = ""
        self.nlp = None
        self.names_ = []
        self.names_tables_ = [] 
        
        self.students = []
        self.teachers = []

        self.n = 10
        self.text_ = ""
        self.tables_ = []
        self.cpus = os.cpu_count()

        try:
            self.nlp = spacy.load("ru_core_news_lg")
            print("Успешная инициализация НЛП модели")

        except OSError:
            print("Error: spaCy model 'ru_core_news_lg' not found.")
            print("Please run: python -m spacy download ru_core_news_lg")

    def __call__(self):
        return self

    def read_document(self, path_: str):
        """
        добавление и считывания файла .docx формата\n
        **path_** - путь к файлу .docx
        """
        if '/' not in path_:
            self.path = f"{pathlib.Path(__file__).parent.resolve()}{self.path}"
        self.path = path_

        try:
            doc = docx.Document(self.path)

            fullText = []

            # чтение текста
            for p in doc.paragraphs:
                fullText.append(p.text)

            self.text_ = '\n'.join(fullText)
            
            # чтение таблиц
            for table in doc.tables:
                table_ = []
                for row in table.rows:
                    row_ = []
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            row_.append(p.text)
                    
                    row_ = ' '.join(row_)
                    table_.append(row_)
                self.tables_.append(table_)
            # print(self.tables_)
                      
        except:
            print("ERROR: Не удалось найти файл!")

    def get_cpus(self) -> int:
        """
        получения количество процессоров, на которое будет распараллелен поиск ФИО студентов и преподавателей
        """
        return self.cpus

    def set_cpus(self, num: int):
        """
        установление количество процессоров, на которое будет распараллелен поиск ФИО студентов и преподавателей
        **num** - количество процессоров для распараллеливания через mpire
        """
        if num <= 0 or num > os.cpu_count():
            print("Некорректное количество процессоров")
            return
        self.cpus = num

    def get_doc_name(self) -> int:
        """
        Получение названия документа .docx
        """
        return self.path.split('/')[-1]
    
    def set_n(self, n: int):
        """
        Установление количества ФИО спанов, которые будут добавляться с конца и начала документа .docx\n
        **n** - количество найденных spaCy НЛП ФИО с начала документа и конца, которые берутся на дальнейшую обработку
        """

        self.n = n

    def search_fio(self, text_: str, st: bool) -> list:
        
        """
        поиск ФИО студентов/преподавателей\n
        **text_** - строка для поиска\n
        **st** - поиск студетов
        """

        pattern = []

        if st:
            pattern = [
            {"LEMMA": { "IN": STUDENTS_IDENTIFICATIONS_WORDS}}
            ] + [
            {"OP": "*"},
            {"TEXT": {"REGEX": "^[А-ЯЁа-яё][а-яё-]+$"}}
        ]
        else:
            pattern = [
            {"LEMMA": { "IN": TEACHERS_IDENTIFICATIONS_WORDS}}
            ] + [
            {"OP": "*"},
            {"TEXT": {"REGEX": "^[А-ЯЁа-яё][а-яё-]+$"}}
        ]


        doc = self.nlp(text_)
        matcher = Matcher(self.nlp.vocab)
        matcher.add("full_names_st", [pattern])
        matches = matcher(doc)

        # print(f"len(matches) = {len(matches)}")

        res_ = []
        # print(doc.ents.__len__())
        
        for match_id, start, end in matches:
            
            span = doc[start:end]
            # print(f"span: {span}")

            res_.append(span)
        
        if len(res_) == 0:
            return

        # print(f"text_: {text_}")
        # print(f"res1: {res_}")

        for i in range(len(res_) - 2, -1, -1):
            if res_[i].text in res_[i + 1].text:
                del res_[i]
        
        print(f"res2: {res_}")

        tokens = []

        for span_ in res_:
            for token in span_:
                if token.ent_type_ == "PER":
                    tokens.append(token)

        if len(tokens) < 1:
            return
        
        print(f"tokens: {tokens}\n")
        return tokens

    def extract_names(self):
        
        """
        Нахождение имен в при помощи NLP библиотеки spaCy
        """

        if not self.nlp or not self.text_:
            print("Не удалось найти имена: модель не была загружена или документ пустой")
            return []

        doc = self.nlp(self.text_)
        # names = [ent.text for ent in doc.ents if ent.label_ == "PER"]
        # self.names_ = names

        # берём только первые 10 найденных имён, дальше - мусор

        # проверка на то, чтобы n всегда было в range для doc.ents.count 
        if self.n > doc.ents.__len__():
            self.n = doc.ents.__len__()
            print(f"doc.ents.__len__() = {doc.ents.__len__()}")

        # поиск ФИО в тексте
        for i in range(self.n):
            if doc.ents[i].label_ == "PER":
                self.names_.append([doc.ents[i].text, [doc.ents[i].start_char, doc.ents[i].end_char]])
            if doc.ents[-i - 1].label == "PER":
                self.names_.append([doc.ents[-i - 1].text, [doc.ents[-i].start_char, doc.ents[-i].end_char]])
        
        # поиск ФИО в таблицах
        # for i in range(len(self.tables_)):
        #     doc_table_ = self.nlp(self.tables_)

        #     for ent in doc_table_.ents:
        #         if ent.label == "PER":

        # генерация паттернов
        patterns_t = [
            {"LOWER": { "IN": TEACHERS_IDENTIFICATIONS_WORDS}}
        ]

        patterns_s = [
            {"LEMMA": { "IN": STUDENTS_IDENTIFICATIONS_WORDS}}
            ] + [
            {"OP": "*"},
            {"TEXT": {"REGEX": "^[А-ЯЁа-яё][а-яё-]+$"}}
        ]


        # поиск работает только для данных из таблиц
        # print(self.tables_[0])

        for i in range(len(self.tables_)):

            text_ = self.tables_[i]

            if isinstance(self.tables_[i], list):
                for j in range(len(text_)):
                    for st in [True, False]:
                        full_name = self.search_fio(text_[j], st)

                        if full_name is not None:
                            if st:
                                self.students.append(full_name)
                            else:
                                self.teachers.append(full_name)

            else:
                for st in [True, False]:
                    full_name = self.search_fio(text_, st)

                    
                    if full_name is not None:
                        if st:
                            self.students.append(full_name)
                        else:
                            self.teachers.append(full_name)

        print("Студенты")
        print(self.students)
        print("\nПреподаватели")
        print(self.teachers)

            # doc = self.nlp(self.tables_[i][0])
            # matcher = Matcher(self.nlp.vocab)
            # matcher.add("full_names_st", [patterns_s])
            # matches = matcher(doc)

            # print(f"len(matches) = {len(matches)}")

            # res_ = []
            # print(doc.ents.__len__())
            
            # for match_id, start, end in matches:
                
            #     span = doc[start:end]
            #     print(f"span: {span}")

            #     res_.append(span.text)
            
            # for i in range(len(res_) - 2, -1, -1):
            #     if res_[i] in res_[i + 1]:
            #         del res_[i]

            # print(res_)

        # print(f"Удалось обнаружить {self.names_.__len__()} предположительных имён в тексте и {self.names_tables_.__len__()} в таблицах")
        # print(self.names_)